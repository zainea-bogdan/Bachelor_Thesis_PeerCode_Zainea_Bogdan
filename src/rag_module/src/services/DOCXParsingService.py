import os
from docx import Document


class DocxParsingService:

    # maximum words per paragraph chunk
    # justified by: Dense Passage Retrieval — Karpukhin et al. 2020
    PARAGRAPH_WINDOW = 200

    # maximum words per section chunk
    # justified by: Lost in the Middle — Liu et al. 2023
    # prevents relevant content from being buried in large chunks
    SECTION_WINDOW = 500

    def parse_docx(self, path: str) -> list[dict]:

        # step 1 — open the document and check its structure
        doc = Document(path)

        # step 2 — collect all heading paragraphs
        headings_found = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                headings_found.append(paragraph)

        # step 3 — decide which strategy to use
        # if the document has no headings, we cannot use
        # heading-based chunking, so we fall back to paragraph chunking
        if len(headings_found) == 0:
            return self._chunk_by_paragraph(path)

        # document has headings, then we use heading-based chunking
        return self._chunk_by_headings(path)

    def _chunk_by_paragraph(self, path: str) -> list[dict]:
        # this method is used when:
        # 1. the document has no heading styles detected
        # 2. as a fallback from _chunk_by_headings

        # open the document
        doc = Document(path)
        filename = os.path.basename(path)

        # results listm here we will hold all chunks produced
        chunks = []

        # chunk counter, which is used to track position in document
        chunk_index = 0

        # rolling word buffer — collects words across paragraphs
        # until we reach PARAGRAPH_WINDOW (200 words)
        current_words = []

        # step through every paragraph in the document one by one
        for paragraph in doc.paragraphs:

            # get the text of this paragraph, remove leading/trailing spaces
            text = paragraph.text.strip()

            # skip this paragraph if it is noise
            # (empty, page number, single word, etc.)
            if self._is_noise(text) == True:
                continue

            # split paragraph into individual words
            words_in_this_paragraph = text.split()

            # add these words to the rolling buffer
            for word in words_in_this_paragraph:
                current_words.append(word)

            # check if the buffer has reached the window size
            if len(current_words) >= self.PARAGRAPH_WINDOW:

                # buffer is full — join words back into a string, cause the split make a list of words
                chunk_text = " ".join(current_words)

                # build the chunk and add it to results
                new_chunk = self._build_chunk(
                    text=chunk_text,
                    filename=filename,
                    chunk_index=chunk_index,
                    strategy="paragraph_chunking"
                )
                chunks.append(new_chunk)

                # increment chunk counter
                chunk_index = chunk_index + 1

                # reset the buffer for the next chunk
                current_words = []

        # end of document reached
        # the last group of words may not have hit 200 words
        # we still need to emit them as a final chunk
        if len(current_words) > 0:
            chunk_text = " ".join(current_words)
            last_chunk = self._build_chunk(
                text=chunk_text,
                filename=filename,
                chunk_index=chunk_index,
                strategy="paragraph_chunking"
            )
            chunks.append(last_chunk)

        return chunks

    def _chunk_by_headings(self, path: str) -> list[dict]:
        # this method groups paragraphs under their heading
        # each heading + its content becomes one chunk
        # if a section is too large (> SECTION_WINDOW),
        # it gets split into 500-word sub-chunks

        doc = Document(path)
        filename = os.path.basename(path)

        chunks = []
        chunk_index = 0

        # the heading text of the section we are currently collecting
        current_heading = None

        # the paragraphs collected under the current heading
        current_paragraphs = []

        # step through every paragraph in the document one by one
        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            # skip noise paragraphs
            if self._is_noise(text) == True:
                continue

            # check if this paragraph is a heading
            if paragraph.style.name.startswith("Heading"):

                # a new heading means the previous section is complete
                # but only flush if we actually collected any paragraphs
                # (handles the case of two headings back to back)
                if len(current_paragraphs) > 0:

                    # flush the previous section into chunks
                    section_chunks = self._flush_section(
                        heading=current_heading,
                        paragraphs=current_paragraphs,
                        filename=filename,
                        chunk_index_start=chunk_index
                    )

                    # add the produced chunks to results
                    for chunk in section_chunks:
                        chunks.append(chunk)

                    # move chunk index forward by however many
                    # chunks this section produced
                    chunk_index = chunk_index + len(section_chunks)

                # start a new section under this heading
                current_heading = text
                current_paragraphs = []

            else:
                # normal paragraph — belongs to the current section
                current_paragraphs.append(text)

        # end of document reached
        # the last section never triggered the heading detector above
        # so we flush it manually here
        if len(current_paragraphs) > 0:
            section_chunks = self._flush_section(
                heading=current_heading,
                paragraphs=current_paragraphs,
                filename=filename,
                chunk_index_start=chunk_index
            )
            for chunk in section_chunks:
                chunks.append(chunk)

        return chunks

    def _flush_section(
        self,
        heading: str | None,
        paragraphs: list[str],
        filename: str,
        chunk_index_start: int
    ) -> list[dict]:
        # this method takes one complete section
        # (heading + its paragraphs) and decides:
        # - if it fits within SECTION_WINDOW → one chunk
        # - if it is too large → split into 500-word sub-chunks

        # step 1 — combine heading and paragraphs into one text block
        if heading is not None:
            # prepend the heading so it travels with the content
            full_text = heading + "\n\n" + " ".join(paragraphs)
        else:
            # section before the first heading has no heading
            full_text = " ".join(paragraphs)

        # step 2 — count total words in this section
        all_words_in_section = full_text.split()
        word_count = len(all_words_in_section)

        # step 3 — decide whether to emit as one chunk or split
        if word_count <= self.SECTION_WINDOW:

            # section fits within limit — emit as one clean chunk
            single_chunk = self._build_chunk(
                text=full_text,
                filename=filename,
                chunk_index=chunk_index_start,
                strategy="heading_chunking"
            )
            return [single_chunk]

        # section is too large — split it
        return self._split_oversized_section(
            full_text=full_text,
            filename=filename,
            chunk_index_start=chunk_index_start
        )

    def _split_oversized_section(
        self,
        full_text: str,
        filename: str,
        chunk_index_start: int
    ) -> list[dict]:
        # this method handles sections that exceed SECTION_WINDOW
        # it slides a 500-word window across the text
        # no overlap needed — section is already semantically
        # bounded by its heading

        # split text into individual words
        all_words = full_text.split()

        chunks = []
        chunk_index = chunk_index_start

        # current position in the word list
        position = 0

        # keep sliding the window until we reach the end
        while position < len(all_words):

            # take the next 500 words
            # if fewer than 500 remain, take whatever is left
            window_words = all_words[position: position + self.SECTION_WINDOW]

            # join back into a string
            chunk_text = " ".join(window_words)

            # build and store this sub-chunk
            new_chunk = self._build_chunk(
                text=chunk_text,
                filename=filename,
                chunk_index=chunk_index,
                strategy="heading_chunking_split_oversized"
            )
            chunks.append(new_chunk)

            # move index and position forward
            chunk_index = chunk_index + 1
            position = position + self.SECTION_WINDOW

        return chunks

    def _build_chunk(
        self,
        text: str,
        filename: str,
        chunk_index: int,
        strategy: str
    ) -> dict:
        # every chunk returned by this service has the same structure
        # this is the contract that IngestionService depends on
        chunk = {
            "text": text,
            "metadata": {
                "source_filename": filename,
                "document_type": "docx",
                "chunk_index": chunk_index,
                "strategy_used": strategy
            }
        }
        return chunk

    @staticmethod
    def _is_noise(text: str) -> bool:
        # empty string — blank line in document
        if not text:
            return True

        # fewer than 3 words — page numbers, labels,
        # single words, footer text
        words = text.split()
        if len(words) < 3:
            return True

        # text is meaningful — keep it
        return False