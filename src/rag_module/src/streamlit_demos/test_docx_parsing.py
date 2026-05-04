import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import streamlit as st
import tempfile
from docx import Document as DocxDocument
from services.DOCXParsingService import DocxParsingService

st.set_page_config(page_title="DocxParsingService — test harness", layout="wide")
st.title("DocxParsingService — test harness")

service = DocxParsingService()

st.sidebar.header("Simulation settings")

mode = st.sidebar.radio(
    "Scenario",
    ["Few DOCX (≤7) — paragraph chunking", "Many DOCX (>7) — heading chunking"]
)

uploaded_file = st.sidebar.file_uploader("Upload a .docx file", type=["docx"])

if mode == "Many DOCX (>7) — heading chunking":
    st.sidebar.markdown("---")
    st.sidebar.markdown("**Simulated Q1/Q3 thresholds**")
    q1 = st.sidebar.slider("Q1 lower threshold (words/section)", 50, 500, 431)
    q3 = st.sidebar.slider("Q3 upper threshold (words/section)", 200, 1200, 637)
    st.sidebar.caption("From corpus analysis: Q1=431.3, Q3=637.38")
else:
    q1 = None
    q3 = None

st.sidebar.markdown("---")
st.sidebar.markdown("**Active thresholds**")
if q1 is None:
    st.sidebar.info("q1 = None → paragraph chunking")
else:
    st.sidebar.success(f"q1 = {q1} | q3 = {q3}")

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    st.subheader("Document structure preview")

    try:
        doc = DocxDocument(tmp_path)
    except Exception as e:
        st.error(f"Failed to open document: {e}")
        os.unlink(tmp_path)
        st.stop()

    headings = [
        p for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    total_words = sum(len(p.text.split()) for p in non_empty)

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total paragraphs", len(non_empty))
    col2.metric("Headings found", len(headings))
    col3.metric("Total words", total_words)
    col4.metric(
        "Avg words/section",
        round(total_words / len(headings), 1) if headings else "N/A"
    )

    if headings:
        avg = round(total_words / len(headings), 2)
        if q1 and q3:
            if avg < q1:
                st.warning(f"avg {avg} < Q1 {q1} → this document would use paragraph chunking even in heading mode (sections too granular)")
            elif avg > q3:
                st.warning(f"avg {avg} > Q3 {q3} → oversized sections will be split further at 500-word windows")
            else:
                st.success(f"avg {avg} is within Q1–Q3 range → clean heading chunking expected")
        else:
            st.info(f"avg words/section = {avg} — heading mode not active")
    else:
        st.warning("No headings detected — paragraph chunking will be used regardless of scenario.")

    st.markdown("---")
    st.subheader("Heading structure")

    if headings:
        for h in headings:
            indent = int(h.style.name.replace("Heading ", "")) - 1 if h.style.name.replace("Heading ", "").isdigit() else 0
            st.markdown(f"{'&nbsp;' * (indent * 4)} **{h.style.name}** — {h.text}", unsafe_allow_html=True)
    else:
        st.write("No headings found in this document.")

    st.markdown("---")

    if st.button("Run parsing", type="primary"):
        with st.spinner("Parsing document..."):
            try:
                chunks = service.parse_docx(tmp_path, q1, q3)
            except Exception as e:
                st.error(f"Parsing failed: {e}")
                os.unlink(tmp_path)
                st.stop()

        st.subheader(f"Results — {len(chunks)} chunks produced")

        strategy_counts = {}
        total_chunk_words = 0
        for c in chunks:
            s = c["metadata"]["strategy_used"]
            strategy_counts[s] = strategy_counts.get(s, 0) + 1
            total_chunk_words += len(c["text"].split())

        res_col1, res_col2, res_col3 = st.columns(3)
        res_col1.metric("Total chunks", len(chunks))
        res_col2.metric("Total words chunked", total_chunk_words)
        res_col3.metric(
            "Avg words/chunk",
            round(total_chunk_words / len(chunks), 1) if chunks else 0
        )

        st.markdown("**Strategy breakdown:**")
        for strategy, count in strategy_counts.items():
            pct = round(count / len(chunks) * 100, 1)
            st.write(f"- `{strategy}`: {count} chunks ({pct}%)")

        st.markdown("---")
        st.subheader("Chunk inspector")

        for i, chunk in enumerate(chunks):
            word_count = len(chunk["text"].split())
            label = f"Chunk {i} — {word_count} words — `{chunk['metadata']['strategy_used']}`"

            with st.expander(label):
                tab1, tab2 = st.tabs(["Text", "Metadata"])

                with tab1:
                    preview = chunk["text"][:800]
                    if len(chunk["text"]) > 800:
                        preview += f"\n\n... [{len(chunk['text']) - 800} more characters]"
                    st.text(preview)

                with tab2:
                    st.json(chunk["metadata"])

    os.unlink(tmp_path)

else:
    st.info("Upload a .docx file from the sidebar to begin.")

    st.subheader("What this tests")

    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("**Scenario 1 — Few DOCX (≤7)**")
        st.markdown("""
        - `q1 = None`, `q3 = None` passed to service
        - Uses `_chunk_by_paragraph`
        - 200-word rolling window
        - Heading structure completely ignored
        - Safe fallback for sparse teacher corpus
        """)

    with col_b:
        st.markdown("**Scenario 2 — Many DOCX (>7)**")
        st.markdown("""
        - Real `q1` and `q3` passed from corpus analysis
        - Uses `_chunk_by_headings`
        - One chunk per heading section
        - Sections above Q3 split at 500-word windows
        - Use sliders to explore threshold sensitivity
        """)

    st.markdown("---")
    st.markdown("**Corpus analysis baseline (from your 26 documents):**")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Q1", "431.3 words/section")
    m2.metric("Q3", "637.38 words/section")
    m3.metric("Std deviation", "964.7")
    m4.metric("Documents without headings", "9 / 26")