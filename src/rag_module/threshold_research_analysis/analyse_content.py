import os
import csv
import statistics
from docx import Document

DOCUMENTS_PATH = os.path.join(os.path.dirname(__file__), "documents")
OUTPUT_CSV = os.path.join(os.path.dirname(__file__), "results.csv")
THRESHOLDS_CSV = os.path.join(os.path.dirname(__file__), "thresholds.csv")


def is_noise(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return True
    if len(text.split()) < 3:
        return True
    return False


def analyse_document(path: str) -> dict:
    try:
        doc = Document(path)
    except Exception as e:
        return {
            "file": os.path.basename(path),
            "total_headings": None,
            "total_words": None,
            "avg_words_per_section": None,
            "recommended_strategy": None,
            "error": str(e)
        }

    headings = [
        p for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]

    all_words = sum(
        len(p.text.split())
        for p in doc.paragraphs
        if not is_noise(p)
    )

    avg_words_per_section = (
        round(all_words / len(headings), 2)
        if headings else None
    )

    return {
        "file": os.path.basename(path),
        "total_headings": len(headings),
        "total_words": all_words,
        "avg_words_per_section": avg_words_per_section,
        "recommended_strategy": None,
        "error": None
    }


def compute_quartiles(values: list[float]) -> tuple[float, float]:
    sorted_values = sorted(values)
    n = len(sorted_values)
    q1 = sorted_values[n // 4]
    q3 = sorted_values[(3 * n) // 4]
    return q1, q3


def classify_strategy(avg: float | None, q1: float, q3: float) -> str:
    if avg is None:
        return "paragraph_chunking (no headings)"
    if avg < q1:
        return "paragraph_chunking (headings too granular)"
    if avg > q3:
        return "heading_chunking + split_oversized"
    return "heading_chunking (optimal)"


def run_analysis():
    files = [
        os.path.join(DOCUMENTS_PATH, f)
        for f in os.listdir(DOCUMENTS_PATH)
        if f.endswith(".docx") and not f.startswith("~$")
    ]

    if not files:
        print("No .docx files found in thresholds_research_analysis/documents/")
        return

    print(f"Found {len(files)} documents. Analysing...")
    results = [analyse_document(f) for f in files]

    failed = [r for r in results if r["error"]]
    if failed:
        print(f"\nWarning: {len(failed)} files failed to parse:")
        for f in failed:
            print(f"  - {f['file']}: {f['error']}")

    valid = [
        r for r in results
        if r["avg_words_per_section"] is not None
    ]

    no_headings = [
        r for r in results
        if r["avg_words_per_section"] is None and not r["error"]
    ]

    if len(valid) < 4:
        print(f"Only {len(valid)} documents with headings found.")
        print("Need at least 4 for reliable quartile computation.")
        return

    avg_values = [r["avg_words_per_section"] for r in valid]
    q1, q3 = compute_quartiles(avg_values)

    for r in results:
        r["recommended_strategy"] = classify_strategy(
            r["avg_words_per_section"], q1, q3
        )

    fieldnames = [
        "file", "total_headings", "total_words",
        "avg_words_per_section", "recommended_strategy", "error"
    ]

    with open(OUTPUT_CSV, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(results)

    with open(THRESHOLDS_CSV, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["metric", "value"])
        writer.writeheader()
        writer.writerows([
            {"metric": "q1_lower_threshold", "value": round(q1, 2)},
            {"metric": "q3_upper_threshold", "value": round(q3, 2)},
            {"metric": "corpus_size_total", "value": len(results)},
            {"metric": "documents_with_headings", "value": len(valid)},
            {"metric": "documents_without_headings", "value": len(no_headings)},
            {"metric": "documents_failed", "value": len(failed)},
            {"metric": "mean_avg_words_per_section",
             "value": round(statistics.mean(avg_values), 2)},
            {"metric": "median_avg_words_per_section",
             "value": round(statistics.median(avg_values), 2)},
            {"metric": "stdev_avg_words_per_section",
             "value": round(statistics.stdev(avg_values), 2)},
            {"metric": "min_avg_words_per_section",
             "value": round(min(avg_values), 2)},
            {"metric": "max_avg_words_per_section",
             "value": round(max(avg_values), 2)},
        ])

    print(f"\nCorpus analysis complete.")
    print(f"  Total documents:       {len(results)}")
    print(f"  With headings:         {len(valid)}")
    print(f"  Without headings:      {len(no_headings)}")
    print(f"  Failed to parse:       {len(failed)}")
    print(f"\n  Q1 threshold:          {round(q1, 2)} words/section")
    print(f"  Q3 threshold:          {round(q3, 2)} words/section")
    print(f"  Mean:                  {round(statistics.mean(avg_values), 2)}")
    print(f"  Median:                {round(statistics.median(avg_values), 2)}")
    print(f"  Std deviation:         {round(statistics.stdev(avg_values), 2)}")
    print(f"\n  Results  -> {OUTPUT_CSV}")
    print(f"  Thresholds -> {THRESHOLDS_CSV}")


if __name__ == "__main__":
    run_analysis()